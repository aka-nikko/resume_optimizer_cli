import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from models.schemas import ResumeOutput
from config import (
    FONT_SIZE,
    PLACEHOLDER_PATTERN,
    EXPERIENCE_PLACEHOLDER,
    PROJECT_PLACEHOLDER,
    MARKDOWN_BOLD_PATTERN,
    REQUIRED_PLACEHOLDERS,
    BULLET_LEFT_INDENT_INCHES,
    BULLET_FIRST_LINE_INDENT_INCHES,
    BULLET_LINE_SPACING,
    BULLET_SPACE_BEFORE_PT,
    BULLET_SPACE_AFTER_PT,
)


class DocxWriter:
    PLACEHOLDER_PATTERN = re.compile(PLACEHOLDER_PATTERN)
    EXPERIENCE_PATTERN = re.compile(EXPERIENCE_PLACEHOLDER)
    PROJECT_PATTERN = re.compile(PROJECT_PLACEHOLDER)
    MARKDOWN_BOLD_PATTERN = re.compile(MARKDOWN_BOLD_PATTERN)

    @staticmethod
    def load_document(path):
        try:
            return Document(path)
        except PackageNotFoundError as exc:
            raise ValueError(f"DOCX file is missing or invalid: {path}") from exc
        except Exception as exc:
            raise RuntimeError(f"Unable to read DOCX file '{path}': {exc}") from exc

    @staticmethod
    def apply_font(paragraph):
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE)

    @staticmethod
    def add_markdown_bold(paragraph, text):
        parts = DocxWriter.MARKDOWN_BOLD_PATTERN.split(text)

        for part in parts:
            if not part:
                continue

            if(part.startswith("**") and part.endswith("**")):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                run = paragraph.add_run(part)
                run.bold = False

            run.font.size = Pt(FONT_SIZE)

    @staticmethod
    def apply_bullet_format(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(BULLET_LEFT_INDENT_INCHES)
        paragraph.paragraph_format.first_line_indent = Inches(BULLET_FIRST_LINE_INDENT_INCHES)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.line_spacing = BULLET_LINE_SPACING
        paragraph.paragraph_format.space_before = Pt(BULLET_SPACE_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER_PT)

    @staticmethod
    def insert_bullet_list(paragraph, bullets):
        if not bullets:
            raise ValueError("Cannot insert an empty bullet list into the template")

        current = paragraph
        current.clear()
        current.add_run("• ")
        DocxWriter.add_markdown_bold(current, bullets[0])
        DocxWriter.apply_bullet_format(current)

        for bullet in bullets[1:]:
            new_p = deepcopy(current._element)
            current._element.addnext(new_p)
            current = Paragraph(new_p,paragraph._parent)
            current.clear()
            current.add_run("• ")
            DocxWriter.add_markdown_bold(current, bullet)
            DocxWriter.apply_bullet_format(current)

    @staticmethod
    def build_skill_section(skills):
        lines = []

        for category, skill_list in ResumeOutput.order_skill_categories(skills).items():
            if not skill_list:
                continue

            lines.append(f"**{category}**: " + ", ".join(skill_list))

        return "\n".join(lines)

    @staticmethod
    def replace_text(paragraph,text):
        paragraph.clear()
        run = paragraph.add_run(text)
        run.font.size = Pt(FONT_SIZE)

    @staticmethod
    def normalize_placeholder_name(name):
        return re.sub(r"[^A-Z0-9]+", "_", name.upper()).strip("_")

    @staticmethod
    def iter_paragraphs(doc):
        for paragraph in doc.paragraphs:
            yield paragraph

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

    @staticmethod
    def get_template_section_counts(template_path):
        doc = DocxWriter.load_document(template_path)
        experience_indexes = set()
        project_indexes = set()

        for paragraph in DocxWriter.iter_paragraphs(doc):
            for match in DocxWriter.PLACEHOLDER_PATTERN.findall(paragraph.text):
                placeholder = DocxWriter.normalize_placeholder_name(match)
                experience_match = DocxWriter.EXPERIENCE_PATTERN.match(placeholder)
                project_match = DocxWriter.PROJECT_PATTERN.match(placeholder)

                if experience_match:
                    experience_indexes.add(int(experience_match.group(1)))

                if project_match:
                    project_indexes.add(int(project_match.group(1)))

        return {
            "experiences": max(experience_indexes, default=0),
            "projects": max(project_indexes, default=0),
        }

    @staticmethod
    def format_field_value(field, value):
        if field == "skills":
            return DocxWriter.build_skill_section(value)

        return value

    @staticmethod
    def build_replacements(resume):
        resume_data = resume.model_dump()
        replacements = {
            "SUMMARY": resume_data["summary"],
            "SKILLS": DocxWriter.format_field_value("skills", resume_data["skills"]),
        }

        for index, bullets in enumerate(resume_data["experiences"], start=1):
            replacements[f"EXPERIENCE_{index}"] = bullets

        for index, bullets in enumerate(resume_data["projects"], start=1):
            replacements[f"PROJECT_{index}"] = bullets

        return replacements

    @staticmethod
    def replace_placeholders(template_path, output_path, resume):
        output = Path(output_path)

        if output.parent:
            output.parent.mkdir(parents=True, exist_ok=True)

        doc = DocxWriter.load_document(template_path)
        replacements = DocxWriter.build_replacements(resume)
        used_placeholders = set()

        for paragraph in DocxWriter.iter_paragraphs(doc):
            matches = DocxWriter.PLACEHOLDER_PATTERN.findall(paragraph.text)
            if not matches:
                continue

            if len(matches) == 1 and DocxWriter.PLACEHOLDER_PATTERN.fullmatch(paragraph.text.strip()):
                placeholder = DocxWriter.normalize_placeholder_name(matches[0])
                value = replacements.get(placeholder)

                if isinstance(value, list):
                    used_placeholders.add(placeholder)
                    DocxWriter.insert_bullet_list(paragraph, value)
                    continue

            new_text = paragraph.text
            for match in matches:
                placeholder = DocxWriter.normalize_placeholder_name(match)
                value = replacements.get(placeholder)
                if value is None:
                    continue

                used_placeholders.add(placeholder)

                if isinstance(value, list):
                    value = "\n".join(value)

                new_text = re.sub(
                    r"\{\{\s*" + re.escape(match) + r"\s*\}\}",
                    lambda _: value,
                    new_text,
                )

            new_text = new_text.strip()

            if new_text == paragraph.text:
                continue

            paragraph.clear()
            DocxWriter.add_markdown_bold(paragraph, new_text)
            paragraph.paragraph_format.line_spacing = BULLET_LINE_SPACING

        missing_required = sorted(set(REQUIRED_PLACEHOLDERS) - used_placeholders)
        if missing_required:
            raise ValueError(
                "Template is missing required placeholder(s): " + ", ".join(missing_required)
            )

        try:
            doc.save(output_path)
        except PermissionError as exc:
            raise PermissionError(
                f"Cannot write output DOCX '{output_path}'. Close the file if it is open and try again."
            ) from exc
        except Exception as exc:
            raise RuntimeError(f"Unable to write output DOCX '{output_path}': {exc}") from exc